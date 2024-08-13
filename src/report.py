# title: 'module report'
# author: 'Elias Albuquerque'
# version: '0.1.0'
# created: '2024-08-10'
# update: '2024-08-10'



import logging
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def report_content(office_object_module, report_path, quote, today, hour, url, screenshot, author):
    """Adiciona o conteúdo do relatório ao documento existente."""

    try:
        doc = Document(report_path)

        # Adiciona o título
        heading_text = "Cotação Atual do Dólar - " + quote + " (" + today + ")"
        heading = doc.add_paragraph(heading_text, style='Heading 1')
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Formatando o texto do título
        run = heading.runs[0]
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 0, 0)  # Preto
        heading.paragraph_format.space_after = Pt(24)

        # Cria um estilo de parágrafo personalizado
        styles = doc.styles
        paragraph_style = styles.add_style('MyParagraphStyle', WD_STYLE_TYPE.PARAGRAPH)
        paragraph_style.font.size = Pt(12)
        paragraph_style.paragraph_format.space_after = Pt(0)
        paragraph_style.paragraph_format.space_before = Pt(0)
        paragraph_style.paragraph_format.line_spacing = Pt(12)

        # Adiciona os parágrafos usando o estilo personalizado
        paragraph = doc.add_paragraph("O dólar está no valor de " + quote, style='MyParagraphStyle')
        paragraph.add_run(", na data " + today + " às " + hour)

        # Adiciona o hyperlink
        paragraph_url = "Valor cotado no site "
        paragraph = doc.add_paragraph(paragraph_url, style='MyParagraphStyle')
        office_object_module.add_hyperlink(paragraph, url, "Banco Central do Brasil.")

        # Adiciona um parágrafo em branco para o espaçamento antes da imagem
        paragraph_blank = doc.add_paragraph()
        paragraph_blank.paragraph_format.space_before = Pt(7)

        # Verifica se a imagem existe antes de tentar adicionar
        if os.path.exists(screenshot):
            picture = doc.add_picture(screenshot, width=Inches(5.88))
        else:
            logging.error(f"Imagem não encontrada: {screenshot}")

        # Insere espaço após a imagem
        paragraph_blank.paragraph_format.space_before = Pt(7)
        doc.add_paragraph("Cotação feita por: " + author, style='MyParagraphStyle')

        # Salva o documento após adicionar todo o conteúdo
        # file_path = os.path.join('reports', 'teste.docx')
        doc.save(report_path)

        logging.info(f"Conteúdo adicionado ao arquivo e salvo em: .\\{report_path}")

    except Exception as e:
        logging.error(f"Erro ao adicionar conteúdo ao documento: {e}")
