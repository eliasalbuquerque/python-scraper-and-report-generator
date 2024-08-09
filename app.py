# title: 'app'
# author: 'Elias Albuquerque'
# version: '0.2.0'
# created: '2024-08-08'
# update: '2024-08-09'


import logging
import json
import datetime
import os
import tempfile
import shutil
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from time import sleep
from src.settings import Settings
from src.website import Website
from src.office import Office


def keep_the_window_open():  # teste
    logging.info('Mantendo site aberto por tempo indeterminado...')
    input('')


def report_content(word, quote, today, hour, url, screenshot, author):
    """Adiciona o conteúdo do relatório ao documento existente."""

    if word.doc is None:
        logging.error("Erro ao criar o documento Word. O relatório não foi gerado.")
        return

    try:
        # Adiciona o título
        heading_text = "Cotação Atual do Dólar - " + quote + " (" + today + ")"
        heading = word.doc.add_paragraph(heading_text, style='Heading 1')
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Formatando o texto do título
        run = heading.runs[0]
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 0, 0)  # Preto
        heading.paragraph_format.space_after = Pt(24)

        # Cria um estilo de parágrafo personalizado
        styles = word.doc.styles
        paragraph_style = styles.add_style('MyParagraphStyle', WD_STYLE_TYPE.PARAGRAPH)
        paragraph_style.font.size = Pt(12)  # Define o tamanho da fonte para 12pt
        paragraph_style.paragraph_format.space_after = Pt(0)  # No Spacing
        paragraph_style.paragraph_format.space_before = Pt(0)  # No Spacing
        paragraph_style.paragraph_format.line_spacing = Pt(12)  # Define o espaçamento entre linhas para 12 pt

        # Adiciona os parágrafos usando o estilo personalizado
        paragraph = word.doc.add_paragraph("O dólar está no valor de " + quote, style='MyParagraphStyle')
        paragraph.add_run(", na data " + today + " às " + hour)
        
        word.doc.add_paragraph("Valor cotado no site " + url, style='MyParagraphStyle')
        word.doc.add_paragraph("Print da cotação atual:", style='MyParagraphStyle')
        
        # Adiciona um parágrafo em branco para o espaçamento antes da imagem
        paragraph_blank = word.doc.add_paragraph()
        paragraph_blank.paragraph_format.space_before = Pt(7)
        
        # Verifica se a imagem existe antes de tentar adicionar
        if os.path.exists(screenshot):
            picture = word.doc.add_picture(screenshot, width=Inches(5.88))
        else:
            logging.error(f"Imagem não encontrada: {screenshot}")
        
        # Insere espaço apos a imagem
        paragraph_blank.paragraph_format.space_before = Pt(7)
        word.doc.add_paragraph("Cotação feita por: " + author, style='MyParagraphStyle')

        # Salva o documento após adicionar todo o conteúdo
        file_path = os.path.join('reports', 'teste.docx')
        word.doc.save(file_path)

        logging.info(f"Conteúdo adicionado ao arquivo e salvo em: .\\{file_path}")

    except Exception as e:
        logging.error(f"Erro ao adicionar conteúdo ao documento: {e}")
    

def get_current_date_time():
    now = datetime.datetime.now()
    return now


def load_config():
    with open('config.json', 'r', encoding='utf8') as file:
        return json.load(file)


def main():
    settings = Settings()
    website = Website(settings)
    word = Office()


    # Criar a pasta temporária
    tempdir = tempfile.mkdtemp()

    try:
        # 0. Carrega as configuracoes e variaveis da aplicacao
        config = load_config()
        url = config['website']['url']
        xp_button_cookie = config['website']['xp_button_cookie']
        xp_quote = config['website']['xp_quote']

        # 1. Acessar o site e extrair o valor da cotacao
        website.access_website(url)
        website.click_on_element(xp_button_cookie)
        website.zoom_out_of_website(78)
        
        quote = website.extract_text_from_element(xp_quote, 'cotação')
        screenshot = website.take_screenshot(tempdir)

        # 2. Criar arquivo Word.docx montar o relatorio
        word.create_document('teste.docx', 'reports')
    
        # 3. Adicionar conteudo no relatorio
        today = get_current_date_time()
        now = get_current_date_time()
        today = now.strftime("%d/%m/%Y")
        hour = now.strftime("%H:%M:%S")
        author = config['office']['author']
        if author == "null":
            author = os.getlogin().capitalize()

        report_content(word, quote, today, hour, url, screenshot, author)

        # * alterar o valor quote para um numero float
        # * adicionar R$ na string do titulo e demais paragrafos
        # * revise as docstrings de todos os modulos
        # * transformar funcao report_content() em um modulo
        # * informar que o modulo report_content é necessario refatorar o codigo caso seja reutilizado em outro projeto

        # 3. Transforme em um PDF
        # - Encontre uma forma de transformar o arquivo word em um arquivo python, usando somente python e transforme o arquivo word em um arquivo pdf
        # 4. Entrega como executável
        # - Transforme o código em um instalador(instruções nas dicas abaixo)

        # keep_the_window_open()
        # os.rmdir(tempdir)

    finally:
        # Remover a pasta temporária no final da execução, se necessário
        shutil.rmtree(tempdir)

if __name__ == '__main__':
    main()
